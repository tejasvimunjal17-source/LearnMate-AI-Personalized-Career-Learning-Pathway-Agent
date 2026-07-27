"""
backend/resume_generator.py
------------------------------
Turns a backend.resume_store.ResumeProfile into an ATS-friendly resume
file (PDF or DOCX), returned as bytes for st.download_button().

Education section format
-------------------------
ResumeProfile.education is a single free-text field (no dedicated
dataclass), populated by frontend/resume_builder.py's dynamic education
cards as blocks like:

    Bachelor of Commerce (Honours) | Pursuing
    Dronacharya Government College
    Gurugram University
    2025 – 2029          Gurugram, Haryana, India

_parse_education_blocks() below segments that free text back into
structured (title, years, institution, location, board, grade) blocks -
anchored on the line that contains a 4-digit year (the "years + location"
line) - and both renderers lay each block out as:

    Degree | Status .......................... 2025 - 2029   (right-aligned)
    Institution ............................... Location       (right-aligned, italic)
    University / Board                                          (bold + italic)

This parsing is self-contained to this module; ResumeProfile and
save_resume() are untouched.

ATS notes
----------
Right-aligned dates require either a table or a right tab stop - a
single-row, borderless 2-column table (PDF) and a right tab stop (DOCX,
no table at all) are used for exactly that purpose. Everything else keeps
the original single-column, bullet-based, no-table layout.

Template-aware rendering (additive)
-------------------------------------
build_resume_pdf()/build_resume_docx() now ALSO read
`backend.resume_templates.get_template(profile.template_id)` to vary
header style (plain/banner/centered-rule/sidebar-bar), fonts, heading
case, divider style, and pill-style skills - plus `profile.accent_color`,
`profile.target_role`, `profile.summary`, and `profile.one_page`. All of
this is read via getattr(profile, "...", default), so this module keeps
working unmodified even if it's ever pointed at an older ResumeProfile
that predates these fields. Both functions gained a new keyword-only
`photo_bytes: bytes | None = None` argument (default None) - every
existing call site (`build_resume_pdf(profile)`) is unaffected; only
callers that opt in by passing photo_bytes=... see the photo embedded,
and only when profile.show_photo is also true.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER
from reportlab.lib import colors as rl_colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, HRFlowable,
)

from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from backend.resume_store import ResumeProfile
from backend.resume_templates import get_template
from backend.logger_setup import get_logger

logger = get_logger(__name__)


class ResumeGenerationError(RuntimeError):
    """Raised when a resume file (PDF or DOCX) fails to generate."""


def _validate_profile(profile: ResumeProfile) -> None:
    if not isinstance(profile, ResumeProfile):
        raise ResumeGenerationError(f"Expected a ResumeProfile instance, got {type(profile).__name__}.")
    if not profile.first_name or not profile.first_name.strip():
        raise ResumeGenerationError("Cannot generate a resume without first_name.")
    if not profile.last_name or not profile.last_name.strip():
        raise ResumeGenerationError("Cannot generate a resume without last_name.")


def _contact_line(profile: ResumeProfile) -> str:
    return profile.email.strip() if profile.email else ""


def _hex(color: str | None, fallback: str = "#2563EB") -> str:
    color = (color or "").strip()
    if not re.match(r"^#[0-9A-Fa-f]{6}$", color):
        return fallback
    return color


# ------------------------------------------------------------------
# Education parsing (free text -> structured blocks) - unchanged
# ------------------------------------------------------------------
# Anchor line: contains a 4-digit year, optionally followed by 2+ spaces
# and a location (e.g. "2025 – 2029          Gurugram, Haryana, India",
# or "2025          Gurugram, Haryana, India" for a single passing year).
_YEAR_LOCATION_RE = re.compile(r"^(?P<years>.*?\d{4}.*?)(?:\s{2,}(?P<location>\S.*))?$")


@dataclass
class _EducationBlock:
    title: str = ""
    years: str = ""
    institution: str = ""
    location: str = ""
    board: str = ""
    grade: str = ""


def _parse_education_blocks(education_text: str) -> list[_EducationBlock]:
    """Segment ResumeProfile.education free text into structured blocks."""
    if not education_text or not education_text.strip():
        return []

    lines = [l for l in education_text.splitlines() if l.strip()]
    raw_blocks: list[list[str]] = []
    buffer: list[str] = []
    i, n = 0, len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()
        buffer.append(line)

        if stripped.lower().startswith("grade:"):
            raw_blocks.append(buffer)
            buffer = []
            i += 1
            continue

        if len(buffer) >= 2 and _YEAR_LOCATION_RE.match(stripped):
            if i + 1 < n and lines[i + 1].strip().lower().startswith("grade:"):
                buffer.append(lines[i + 1])
                i += 2
            else:
                i += 1
            raw_blocks.append(buffer)
            buffer = []
            continue

        i += 1

    if buffer:
        raw_blocks.append(buffer)  # leftover lines that didn't match the pattern

    blocks: list[_EducationBlock] = []
    for raw in raw_blocks:
        block = _EducationBlock()
        remaining = list(raw)

        if remaining and remaining[-1].strip().lower().startswith("grade:"):
            block.grade = remaining[-1].split(":", 1)[-1].strip()
            remaining = remaining[:-1]

        if remaining:
            m = _YEAR_LOCATION_RE.match(remaining[-1].strip())
            if m:
                block.years = (m.group("years") or "").strip()
                block.location = (m.group("location") or "").strip()
                remaining = remaining[:-1]

        if remaining:
            block.title = remaining[0].strip()
        if len(remaining) >= 2:
            block.institution = remaining[1].strip()
        if len(remaining) >= 3:
            block.board = remaining[2].strip()

        blocks.append(block)

    return blocks


# ------------------------------------------------------------------
# PDF (reportlab)
# ------------------------------------------------------------------
def _add_education_pdf(story: list, blocks: list[_EducationBlock], heading_style, content_width: float,
                        gap: float = 6) -> None:
    title_style = ParagraphStyle("EduTitle", fontName="Helvetica-Bold", fontSize=10, leading=14, alignment=TA_LEFT)
    year_style = ParagraphStyle("EduYear", fontName="Helvetica", fontSize=10, leading=14, alignment=TA_RIGHT)
    inst_style = ParagraphStyle("EduInst", fontName="Helvetica", fontSize=10, leading=14, alignment=TA_LEFT)
    loc_style = ParagraphStyle("EduLoc", fontName="Helvetica-Oblique", fontSize=10, leading=14, alignment=TA_RIGHT)
    board_style = ParagraphStyle("EduBoard", fontName="Helvetica-BoldOblique", fontSize=10, leading=14, alignment=TA_LEFT)
    grade_style = ParagraphStyle("EduGrade", fontName="Helvetica", fontSize=9, leading=12,
                                  alignment=TA_LEFT, textColor="#444444")

    story.append(Paragraph("EDUCATION", heading_style))
    left_w = content_width * 0.62
    right_w = content_width - left_w

    for idx, b in enumerate(blocks):
        rows = []
        if b.title or b.years:
            rows.append([Paragraph(b.title, title_style), Paragraph(b.years, year_style)])
        if b.institution or b.location:
            rows.append([Paragraph(b.institution, inst_style), Paragraph(b.location, loc_style)])

        if rows:
            table = Table(rows, colWidths=[left_w, right_w])
            table.setStyle(TableStyle([
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 1),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(table)

        if b.board:
            story.append(Paragraph(b.board, board_style))
        if b.grade:
            story.append(Paragraph(f"Grade: {b.grade}", grade_style))

        if idx < len(blocks) - 1:
            story.append(Spacer(1, gap))


def build_resume_pdf(
    profile: ResumeProfile,
    *,
    template_id: str | None = None,
    accent_color: str | None = None,
    photo_bytes: bytes | None = None,
) -> bytes:
    """Generate an ATS-friendly PDF resume, styled per the selected template.

    Args:
        profile: The ResumeProfile to render.
        template_id: Overrides profile.template_id if given.
        accent_color: Overrides profile.accent_color if given (hex, e.g. "#2563EB").
        photo_bytes: Optional headshot image bytes, embedded only if
            profile.show_photo is True.

    Returns:
        The generated PDF file's raw bytes.

    Raises:
        ResumeGenerationError: if `profile` is invalid or PDF rendering fails.
    """
    _validate_profile(profile)

    template = get_template(template_id or getattr(profile, "template_id", None))
    style = template.style
    accent_hex = _hex(accent_color or getattr(profile, "accent_color", None))
    accent = rl_colors.HexColor(accent_hex)
    ink = rl_colors.HexColor("#1a1a1a")
    muted = rl_colors.HexColor("#4b5563")
    one_page = getattr(profile, "one_page", True)
    if one_page is None:
        one_page = True

    sp_before = 8 if one_page else 12
    sp_after = 3 if one_page else 5
    edu_gap = 5 if one_page else 8
    top_margin = 0.5 if one_page else 0.65

    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, pagesize=LETTER,
            topMargin=top_margin * inch, bottomMargin=0.6 * inch,
            leftMargin=0.7 * inch, rightMargin=0.7 * inch,
            title=f"{profile.full_name} - Resume",
        )
        styles = getSampleStyleSheet()

        name_align = TA_CENTER if style.name_align == "center" else TA_LEFT
        header_is_banner = style.header_style == "banner"
        name_color = rl_colors.white if header_is_banner else (accent if style.header_style == "sidebar-bar" else ink)

        name_style = ParagraphStyle(
            "NameStyle", parent=styles["Title"], fontName=style.heading_font,
            fontSize=18, alignment=name_align, spaceAfter=2, textColor=name_color,
        )
        tagline_style = ParagraphStyle(
            "TaglineStyle", parent=styles["Normal"], fontName=style.body_font, fontSize=10.5,
            alignment=name_align,
            textColor=(rl_colors.HexColor("#e5e7eb") if header_is_banner else accent), spaceAfter=2,
        )
        contact_style = ParagraphStyle(
            "ContactStyle", parent=styles["Normal"], fontName=style.body_font, fontSize=9.5,
            alignment=name_align,
            textColor=(rl_colors.HexColor("#f3f4f6") if header_is_banner else muted), spaceAfter=10,
        )
        heading_style = ParagraphStyle(
            "SectionHeading", parent=styles["Heading2"], fontName=style.heading_font,
            fontSize=11.5, spaceBefore=sp_before, spaceAfter=4,
            textColor=(accent if style.header_style in ("banner", "sidebar-bar") else ink),
        )
        body_style = ParagraphStyle(
            "Body", parent=styles["Normal"], fontName=style.body_font, fontSize=10, leading=14, spaceAfter=sp_after,
        )
        bullet_style = ParagraphStyle(
            "Bullet", parent=body_style, leftIndent=12, bulletIndent=0,
        )
        summary_style = ParagraphStyle(
            "Summary", parent=body_style, fontName=style.body_font, spaceAfter=sp_after,
        )
        skills_style = ParagraphStyle("Skills", parent=body_style, textColor=ink)

        name_text = profile.full_name or "Your Name"
        tagline_text = (getattr(profile, "target_role", "") or "").strip()
        contact = _contact_line(profile)

        header_flowables: list = [Paragraph(name_text, name_style)]
        if tagline_text:
            header_flowables.append(Paragraph(tagline_text, tagline_style))
        header_flowables.append(Paragraph(contact, contact_style) if contact else Spacer(1, 6))

        photo_flowable = None
        if getattr(profile, "show_photo", False) and photo_bytes:
            try:
                photo_flowable = Image(io.BytesIO(photo_bytes), width=0.85 * inch, height=0.85 * inch)
            except Exception:  # noqa: BLE001 - a bad image should never break resume generation
                photo_flowable = None

        story: list = []

        if style.header_style == "sidebar-bar":
            story.append(HRFlowable(width="100%", thickness=4, color=accent, spaceAfter=8, lineCap="round"))

        if photo_flowable is not None:
            header_table = Table(
                [[header_flowables, photo_flowable]],
                colWidths=[doc.width - 1.0 * inch, 1.0 * inch],
            )
            header_table.setStyle(TableStyle([
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (1, 0), (1, 0), "RIGHT"),
            ]))
            header_content = header_table
        else:
            header_content = header_flowables

        if header_is_banner:
            banner_table = Table([[header_content]], colWidths=[doc.width])
            banner_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), accent),
                ("LEFTPADDING", (0, 0), (-1, -1), 14),
                ("RIGHTPADDING", (0, 0), (-1, -1), 14),
                ("TOPPADDING", (0, 0), (-1, -1), 12),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
            ]))
            story.append(banner_table)
            story.append(Spacer(1, 10))
        else:
            if isinstance(header_content, list):
                story.extend(header_content)
            else:
                story.append(header_content)
            if style.header_style == "centered-rule":
                story.append(Spacer(1, 4))
                story.append(HRFlowable(width="100%", thickness=1.4, color=accent, spaceAfter=2))
                story.append(HRFlowable(width="100%", thickness=0.5, color=ink, spaceAfter=8))

        def add_heading(text: str) -> None:
            label = text.upper() if style.uppercase_headings else text
            story.append(Paragraph(label, heading_style))
            if style.divider:
                story.append(HRFlowable(width="100%", thickness=0.6,
                                         color=rl_colors.HexColor("#d1d5db"), spaceAfter=4))

        summary_text = (getattr(profile, "summary", "") or "").strip()
        if summary_text:
            add_heading("Professional Summary")
            story.append(Paragraph(summary_text, summary_style))

        if profile.education:
            education_blocks = _parse_education_blocks(profile.education)
            if education_blocks:
                _add_education_pdf(story, education_blocks, heading_style, doc.width, edu_gap)
            else:
                add_heading("Education")
                for line in profile.education.splitlines():
                    if line.strip():
                        story.append(Paragraph(f"- {line.strip()}", bullet_style))

        if profile.skills:
            add_heading("Skills")
            if style.pill_skills:
                skills_html = "&nbsp;&nbsp;".join(
                    f'<font color="{accent_hex}">\u25CF</font> {s}' for s in profile.skills
                )
                story.append(Paragraph(skills_html, skills_style))
            else:
                story.append(Paragraph(", ".join(profile.skills), body_style))

        if profile.internships:
            add_heading("Internship Experience")
            for i in profile.internships:
                header = f"<b>{i.role}</b>"
                if i.company:
                    header += f" - {i.company}"
                if i.duration:
                    header += f" ({i.duration})"
                story.append(Paragraph(header, body_style))
                if i.description:
                    story.append(Paragraph(f"- {i.description}", bullet_style))

        if profile.projects:
            add_heading("Projects")
            for p in profile.projects:
                title = p.title + (f" ({p.tech_stack})" if p.tech_stack else "")
                story.append(Paragraph(f"<b>{title}</b>", body_style))
                if p.description:
                    story.append(Paragraph(f"- {p.description}", bullet_style))

        if profile.certificates:
            add_heading("Certifications")
            for c in profile.certificates:
                label = c.name
                if c.issuer:
                    label += f" - {c.issuer}"
                if c.year:
                    label += f" ({c.year})"
                story.append(Paragraph(f"- {label}", bullet_style))

        if profile.achievements:
            add_heading("Achievements")
            for line in profile.achievements.splitlines():
                if line.strip():
                    story.append(Paragraph(f"- {line.strip()}", bullet_style))

        if profile.hobbies:
            add_heading("Hobbies")
            story.append(Paragraph(", ".join(profile.hobbies), body_style))

        story.append(Spacer(1, 4))
        doc.build(story)
        return buffer.getvalue()

    except ResumeGenerationError:
        raise
    except Exception as exc:  # noqa: BLE001
        logger.exception("Failed to generate PDF resume for %s", getattr(profile, "email", "unknown"))
        raise ResumeGenerationError(f"Could not generate PDF resume: {exc}") from exc


# ------------------------------------------------------------------
# DOCX (python-docx)
# ------------------------------------------------------------------
def _set_paragraph_shading(paragraph, hex_color: str) -> None:
    """Apply a solid background shade to a paragraph (used for banner headers)."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    pPr.append(shd)


def _set_paragraph_border(paragraph, hex_color: str, size: int = 6, position: str = "bottom") -> None:
    """Add a single bottom (or top) border rule under a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    edge = OxmlElement(f"w:{position}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), "4")
    edge.set(qn("w:color"), hex_color.lstrip("#"))
    pBdr.append(edge)
    pPr.append(pBdr)


def _add_education_docx(doc: Document, blocks: list[_EducationBlock], add_heading) -> None:
    section = doc.sections[0]
    usable_width_in = Emu(
        section.page_width.emu - section.left_margin.emu - section.right_margin.emu
    ).inches

    add_heading("Education")

    for idx, b in enumerate(blocks):
        if b.title or b.years:
            p = doc.add_paragraph()
            p.paragraph_format.tab_stops.add_tab_stop(Inches(usable_width_in), WD_TAB_ALIGNMENT.RIGHT)
            title_run = p.add_run(b.title)
            title_run.bold = True
            if b.years:
                p.add_run(f"\t{b.years}")

        if b.institution or b.location:
            p2 = doc.add_paragraph()
            p2.paragraph_format.tab_stops.add_tab_stop(Inches(usable_width_in), WD_TAB_ALIGNMENT.RIGHT)
            p2.add_run(b.institution)
            if b.location:
                loc_run = p2.add_run(f"\t{b.location}")
                loc_run.italic = True

        if b.board:
            p3 = doc.add_paragraph()
            board_run = p3.add_run(b.board)
            board_run.bold = True
            board_run.italic = True

        if b.grade:
            grade_p = doc.add_paragraph(f"Grade: {b.grade}")
            if grade_p.runs:
                grade_p.runs[0].font.size = Pt(9)

        if idx < len(blocks) - 1:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(2)


def build_resume_docx(
    profile: ResumeProfile,
    *,
    template_id: str | None = None,
    accent_color: str | None = None,
    photo_bytes: bytes | None = None,
) -> bytes:
    """Generate an ATS-friendly DOCX resume, styled per the selected template.

    Args mirror build_resume_pdf(). Returns raw DOCX bytes.

    Raises:
        ResumeGenerationError: if `profile` is invalid or DOCX rendering fails.
    """
    _validate_profile(profile)

    template = get_template(template_id or getattr(profile, "template_id", None))
    style = template.style
    accent_hex = _hex(accent_color or getattr(profile, "accent_color", None)).lstrip("#")
    accent_rgb = RGBColor(int(accent_hex[0:2], 16), int(accent_hex[2:4], 16), int(accent_hex[4:6], 16))
    one_page = getattr(profile, "one_page", True)
    if one_page is None:
        one_page = True
    heading_space_before = Pt(8 if one_page else 12)

    try:
        doc = Document()

        base_style = doc.styles["Normal"]
        base_style.font.name = style.docx_font
        base_style.font.size = Pt(10.5 if one_page else 11)

        for section in doc.sections:
            section.top_margin = section.bottom_margin = Pt(32 if one_page else 40)
            section.left_margin = section.right_margin = Pt(50)

        align_center = style.name_align == "center"
        header_is_banner = style.header_style == "banner"

        if style.header_style == "sidebar-bar":
            bar_p = doc.add_paragraph()
            _set_paragraph_shading(bar_p, accent_hex)
            bar_p.paragraph_format.space_after = Pt(2)
            run = bar_p.add_run(" ")
            run.font.size = Pt(2)

        name_p = doc.add_paragraph()
        if align_center:
            name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if header_is_banner:
            _set_paragraph_shading(name_p, accent_hex)
            name_p.paragraph_format.space_before = Pt(10)
        name_run = name_p.add_run(profile.full_name or "Your Name")
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_run.font.name = style.docx_font
        if header_is_banner:
            name_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif style.header_style == "sidebar-bar":
            name_run.font.color.rgb = accent_rgb

        tagline_text = (getattr(profile, "target_role", "") or "").strip()
        if tagline_text:
            tagline_p = doc.add_paragraph()
            if align_center:
                tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if header_is_banner:
                _set_paragraph_shading(tagline_p, accent_hex)
            tagline_run = tagline_p.add_run(tagline_text)
            tagline_run.font.size = Pt(11)
            tagline_run.font.color.rgb = (RGBColor(0xF3, 0xF4, 0xF6) if header_is_banner else accent_rgb)

        contact = _contact_line(profile)
        contact_p = doc.add_paragraph()
        if align_center:
            contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if header_is_banner:
            _set_paragraph_shading(contact_p, accent_hex)
            contact_p.paragraph_format.space_after = Pt(10)
        if contact:
            contact_run = contact_p.add_run(contact)
            contact_run.font.size = Pt(9.5)
            if header_is_banner:
                contact_run.font.color.rgb = RGBColor(0xF3, 0xF4, 0xF6)

        if getattr(profile, "show_photo", False) and photo_bytes:
            try:
                doc.add_picture(io.BytesIO(photo_bytes), width=Inches(0.9))
            except Exception:  # noqa: BLE001
                pass

        if not header_is_banner:
            rule_p = doc.add_paragraph()
            _set_paragraph_border(rule_p, accent_hex if style.header_style == "centered-rule" else "D1D5DB",
                                   size=(10 if style.header_style == "centered-rule" else 6))
            rule_p.paragraph_format.space_after = Pt(6)

        def add_heading(text: str) -> None:
            label = text.upper() if style.uppercase_headings else text
            h = doc.add_paragraph()
            run = h.add_run(label)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = style.docx_font
            if style.header_style in ("banner", "sidebar-bar"):
                run.font.color.rgb = accent_rgb
            h.paragraph_format.space_before = heading_space_before
            h.paragraph_format.space_after = Pt(2)
            if style.divider:
                _set_paragraph_border(h, "D1D5DB", size=4)

        def add_bullet(text: str) -> None:
            doc.add_paragraph(f"- {text}")

        summary_text = (getattr(profile, "summary", "") or "").strip()
        if summary_text:
            add_heading("Professional Summary")
            doc.add_paragraph(summary_text)

        if profile.education:
            education_blocks = _parse_education_blocks(profile.education)
            if education_blocks:
                _add_education_docx(doc, education_blocks, add_heading)
            else:
                add_heading("Education")
                for line in profile.education.splitlines():
                    if line.strip():
                        add_bullet(line.strip())

        if profile.skills:
            add_heading("Skills")
            if style.pill_skills:
                p = doc.add_paragraph()
                for idx, skill in enumerate(profile.skills):
                    if idx > 0:
                        p.add_run("    ")
                    dot = p.add_run("\u25CF ")
                    dot.font.color.rgb = accent_rgb
                    p.add_run(skill)
            else:
                doc.add_paragraph(", ".join(profile.skills))

        if profile.internships:
            add_heading("Internship Experience")
            for i in profile.internships:
                p = doc.add_paragraph()
                header = i.role
                if i.company:
                    header += f" - {i.company}"
                if i.duration:
                    header += f" ({i.duration})"
                p.add_run(header).bold = True
                if i.description:
                    add_bullet(i.description)

        if profile.projects:
            add_heading("Projects")
            for proj in profile.projects:
                p = doc.add_paragraph()
                title = proj.title + (f" ({proj.tech_stack})" if proj.tech_stack else "")
                p.add_run(title).bold = True
                if proj.description:
                    add_bullet(proj.description)

        if profile.certificates:
            add_heading("Certifications")
            for c in profile.certificates:
                label = c.name
                if c.issuer:
                    label += f" - {c.issuer}"
                if c.year:
                    label += f" ({c.year})"
                add_bullet(label)

        if profile.achievements:
            add_heading("Achievements")
            for line in profile.achievements.splitlines():
                if line.strip():
                    add_bullet(line.strip())

        if profile.hobbies:
            add_heading("Hobbies")
            doc.add_paragraph(", ".join(profile.hobbies))

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    except ResumeGenerationError:
        raise
    except Exception as exc:  # noqa: BLE001
        logger.exception("Failed to generate DOCX resume for %s", getattr(profile, "email", "unknown"))
        raise ResumeGenerationError(f"Could not generate DOCX resume: {exc}") from exc